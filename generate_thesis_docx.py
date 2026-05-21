from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_XIAOWU = 9
FONT_WU = 10.5
FONT_XIAOSI = 12
FONT_SI = 14
FONT_SAN = 16
FONT_ER = 22
INDENT_2CH_CM = 0.74
BOOKMARK_COUNTER = 1


ROOT = Path(__file__).resolve().parent
SOURCE_CANDIDATES = [
    ROOT / "main_add_ai_search_transfer.tex",
    ROOT / "图片" / "main_fixed.tex",
    ROOT / "图片" / "main.tex",
]
OUTPUT_DOCX = ROOT / "基于Python的在线网址收藏管理平台的设计与实现（可编辑）.docx"
IMAGE_ROOT = ROOT / "图片"

TITLE_CN = "基于Python的在线网址收藏管理平台的设计与实现"
TITLE_EN = "DESIGN AND IMPLEMENTATION OF AN ONLINE URL COLLECTION MANAGEMENT PLATFORM BASED ON PYTHON"

AUTHOR = "杨永坤"
STUDENT_ID = "202209080105"
GRADE = "2022"
MAJOR = "软件工程"
COLLEGE = "信息科学与工程"
SUPERVISOR = "王永芳"

ORIGINALITY_TEXT = (
    "郑重声明：所呈交的学位论文《基于Python的在线网址收藏管理平台的设计与实现》，是在指导教师指导下独立完成的研究成果。"
    "除文中已经注明引用的内容外，本论文不包含任何他人已经发表或撰写过的研究成果，也不包含为获得其他教育机构学位或证书而使用过的材料。"
    "对本文研究做出贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律后果，并愿意承担由此引起的全部责任。"
)

AUTHORIZATION_TEXT = (
    "本学位论文作者完全了解学校有关保留、使用学位论文的规定，同意学校保留并向国家有关部门或机构送交论文的复印件和电子版，"
    "允许论文被查阅和借阅。本人授权临沂大学将本学位论文的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存和汇编本学位论文。"
)

ABSTRACT_CN = [
    "在互联网信息规模快速增长的背景下，用户对网页收藏、分类整理、跨设备访问以及内容检索的需求不断增强。传统浏览器收藏夹主要承担静态保存功能，缺乏系统化分类、公开共享、行为统计和智能辅助等能力，难以满足用户在个人知识管理与轻社交传播场景下的实际使用需求。针对上述问题，本文设计并实现了一套基于 Python 的在线网址收藏管理平台，以提高网址资源管理效率、增强内容组织能力，并探索 AI 技术在网址收藏服务领域中的应用价值。",
    "本系统采用前后端分离架构，前端基于 Vue 3、Vue Router、Pinia、Element Plus 和 ECharts 实现用户端与管理端界面；后端基于 Python Flask、Flask-SQLAlchemy 与 Flask-JWT-Extended 构建 REST 风格接口，数据库采用 MySQL。围绕实际应用需求，系统实现了用户注册登录、个人书签管理、树形分类管理、公开书签展示、热门榜单、搜索功能、访问统计、用户关注与点赞、后台审核治理、日志记录、AI 设置以及双向会话式反馈等功能模块。针对传统反馈系统交互单向、进度不可追踪的问题，本文进一步设计了由反馈主体与消息记录组成的会话式反馈模型，形成用户与管理员之间的闭环沟通机制。",
    "系统实现结果表明，该平台能够较好地完成网址采集、组织、共享、统计与智能辅助分析等核心任务，兼顾了课程设计中的工程完整性与后续扩展能力。测试结果表明，各主要功能模块运行稳定，认证流程、书签管理、反馈会话以及数据库持久化能力均能够满足预期要求。本文所实现的平台为在线网址收藏管理系统的构建提供了一种可落地的设计方案，也为后续在推荐算法、全文检索和消息推送等方向的拓展奠定了基础。"
]

KEYWORDS_CN = "在线网址收藏；Python；Flask；Vue 3；前后端分离；会话式反馈"

ABSTRACT_EN = [
    "With the rapid growth of Internet information, users increasingly demand efficient web URL collection, hierarchical classification, cross-device access, content retrieval, and intelligent assistance. Traditional browser bookmarks mainly provide static storage functions and generally lack systematic organization, public sharing, behavioral analytics, and intelligent support. Therefore, they are insufficient for personal knowledge management and lightweight social sharing scenarios. To address these issues, this paper designs and implements an online URL collection management platform based on Python, aiming to improve the efficiency of URL organization and explore the value of AI-enhanced web information services.",
    "The system adopts a front-end and back-end separated architecture. The front end is built with Vue 3, Vue Router, Pinia, Element Plus, and ECharts to provide both user-side and administrator-side interfaces. The back end is developed with Python Flask, Flask-SQLAlchemy, and Flask-JWT-Extended to provide REST-style APIs. MySQL is used as the database of the system. According to the practical requirements of the project, the system implements user registration and login, personal bookmark management, tree-structured category management, public bookmark display, ranking list, global search, access statistics, user follow and like functions, administrator review and governance, operation logging, AI configuration, and a bidirectional conversational feedback module.",
    "Experimental use and functional testing show that the system can effectively support URL collection, organization, sharing, statistics, and intelligent assistance while maintaining good extensibility and engineering completeness. The test results indicate that the authentication flow, bookmark management, feedback conversation, and database persistence functions all meet the expected requirements. The implemented system provides a feasible solution for online URL collection platforms and lays a foundation for future work in recommendation algorithms, full-text retrieval, and real-time notification services."
]

KEYWORDS_EN = "online URL collection; Python; Flask; Vue 3; front-end and back-end separation; conversational feedback"


def set_run_font(run, size: float = FONT_XIAOSI, bold: bool = False, chinese: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)
    run.bold = bold


def style_paragraph(
    paragraph,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent_cm: float | None = None,
    line_spacing: float = 1.25,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)


def add_text(
    doc: Document,
    text: str,
    *,
    size: float = FONT_XIAOSI,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent_cm: float | None = INDENT_2CH_CM,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    chinese: str = "宋体",
    latin: str = "Times New Roman",
) -> None:
    p = doc.add_paragraph()
    style_paragraph(
        p,
        align=align,
        first_line_indent_cm=first_line_indent_cm,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
    )
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, chinese=chinese, latin=latin)


def add_heading_text(doc: Document, text: str, level: int, bookmark_name: str | None = None) -> None:
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, space_before_pt=12, space_after_pt=12)
        run = p.add_run(text)
        set_run_font(run, size=FONT_SAN, bold=True, chinese="黑体")
    elif level == 2:
        style_paragraph(p, first_line_indent_cm=None, space_before_pt=6, space_after_pt=6)
        run = p.add_run(text)
        set_run_font(run, size=FONT_SI, bold=True, chinese="黑体")
    else:
        style_paragraph(p, first_line_indent_cm=None, space_before_pt=6, space_after_pt=6)
        run = p.add_run(text)
        set_run_font(run, size=FONT_XIAOSI, bold=True, chinese="黑体")
    if bookmark_name:
        add_bookmark(p, bookmark_name)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_bookmark(paragraph, name: str) -> None:
    global BOOKMARK_COUNTER
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(BOOKMARK_COUNTER))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(BOOKMARK_COUNTER))
    BOOKMARK_COUNTER += 1
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")


def add_pageref_field(paragraph, bookmark_name: str, default_text: str = "1") -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'PAGEREF {bookmark_name}')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = default_text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def start_new_section(doc: Document, header_text: str | None = None) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ""
    if header_text:
        p = section.header.paragraphs[0]
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, line_spacing=1.0)
        r = p.add_run(header_text)
        set_run_font(r, size=9, chinese="宋体")
        add_bottom_border(p)


def add_placeholder(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, space_before_pt=6, space_after_pt=6)
    run = p.add_run(f"【此处插入{caption}】")
    set_run_font(run, size=FONT_WU, chinese="楷体")


def add_image_or_placeholder(doc: Document, image_name: str, caption: str, width_cm: float = 14.5) -> None:
    image_path = IMAGE_ROOT / image_name
    if image_path.exists():
        p = doc.add_paragraph()
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, space_before_pt=6, space_after_pt=6)
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        add_placeholder(doc, caption)
    add_text(doc, caption, size=FONT_WU, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="楷体")


def add_field_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, first_line_indent_cm=None)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录将在 Word 中更新后显示"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def build_outline(tex: str) -> list[tuple[int, str, str]]:
    body = extract_block(tex, r"\section{绪论}", r"\begin{thebibliography}{99}")
    body = r"\section{绪论}" + body
    outline: list[tuple[int, str, str]] = []
    sec = sub = subsub = 0
    for raw in body.splitlines():
        line = raw.strip()
        if line.startswith(r"\section{"):
            sec += 1
            sub = 0
            subsub = 0
            title = clean_inline(re.search(r"\\section\{(.*)\}", line).group(1))
            outline.append((1, f"{sec} {title}", f"toc_{sec}"))
        elif line.startswith(r"\subsection{"):
            sub += 1
            subsub = 0
            title = clean_inline(re.search(r"\\subsection\{(.*)\}", line).group(1))
            outline.append((2, f"{sec}.{sub} {title}", f"toc_{sec}_{sub}"))
        elif line.startswith(r"\subsubsection{"):
            subsub += 1
            title = clean_inline(re.search(r"\\subsubsection\{(.*)\}", line).group(1))
            outline.append((3, f"{sec}.{sub}.{subsub} {title}", f"toc_{sec}_{sub}_{subsub}"))
    return outline


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\\texttt\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\makecell\[l\]\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\underline\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\checkedbox", "■", text)
    text = re.sub(r"\\quad", " ", text)
    text = re.sub(r"\\hspace\{[^{}]+\}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?", "", text)
    text = text.replace("`", "")
    text = text.replace("{", "").replace("}", "")
    text = text.replace("~", " ")
    text = text.replace(r"\_", "_").replace(r"\&", "&").replace(r"\%", "%")
    text = text.replace("pendingreplied", "pending/replied")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_rows(tabular_text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in tabular_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith(r"\toprule") or line.startswith(r"\midrule") or line.startswith(r"\bottomrule"):
            continue
        if line.startswith(r"\endfirsthead") or line.startswith(r"\endhead") or line.startswith(r"\endfoot") or line.startswith(r"\endlastfoot"):
            continue
        if line.startswith(r"\caption") or line.startswith(r"\label"):
            continue
        line = line.rstrip("\\").strip()
        if not line:
            continue
        cells = [clean_inline(c) for c in line.split("&")]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    # Some longtable blocks repeat the header after \endfirsthead; drop repeated header rows.
    if len(rows) >= 2:
        deduped = [rows[0]]
        for row in rows[1:]:
            if row != rows[0]:
                deduped.append(row)
        rows = deduped
    return rows


def add_docx_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    if not rows:
        add_placeholder(doc, caption)
        return
    add_text(doc, caption, size=FONT_SI, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="宋体")
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            style_paragraph(p, first_line_indent_cm=None)
            run = p.add_run(cell_text)
            set_run_font(run, size=FONT_WU if i else FONT_XIAOSI, bold=(i == 0))
    doc.add_paragraph()


def extract_block(text: str, start: str, end: str) -> str:
    start_idx = text.find(start)
    if start_idx == -1:
        return ""
    start_idx += len(start)
    end_idx = text.find(end, start_idx)
    if end_idx == -1:
        return text[start_idx:]
    return text[start_idx:end_idx]


def parse_table_block(block: str) -> tuple[str, list[list[str]]]:
    cap_match = re.search(r"\\caption\{([^{}]+)\}", block, re.S)
    caption = clean_inline(cap_match.group(1)) if cap_match else "表格占位"
    tabular_match = re.search(r"\\begin\{tabular\}\{[^\n]*\}\s*(.*?)\\end\{tabular\}", block, re.S)
    if not tabular_match:
        tabular_match = re.search(r"\\begin\{longtable\}\{[^\n]*\}\s*(.*?)\\end\{longtable\}", block, re.S)
    rows = parse_rows(tabular_match.group(1)) if tabular_match else []
    return caption, rows


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    add_text(doc, "2026届", size=FONT_XIAOSI, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=None, chinese="黑体")

    p = doc.add_paragraph()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent_cm=None)
    run = p.add_run("分类号：TP311\n单位代码：10452")
    set_run_font(run, size=FONT_XIAOSI, bold=False, chinese="黑体")

    add_text(doc, "【此处插入临沂大学校名字样图片】", size=18, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="楷体")
    doc.add_paragraph()
    add_text(doc, "毕业论文（设计）", size=FONT_ER, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="黑体")
    doc.add_paragraph()
    add_text(doc, TITLE_CN, size=FONT_SAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="黑体")
    doc.add_paragraph()

    info = [
        ("姓    名", AUTHOR),
        ("学    号", STUDENT_ID),
        ("年    级", GRADE),
        ("专    业", MAJOR),
        ("学    院", COLLEGE),
        ("指导教师", SUPERVISOR),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None)
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, size=14, chinese="黑体")

    doc.add_paragraph()
    doc.add_paragraph()
    add_text(doc, "2026年4月", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="黑体")


def add_originality_and_authorization(doc: Document) -> None:
    add_page_break(doc)
    add_heading_text(doc, "学位论文原创性声明", 1)
    add_text(doc, ORIGINALITY_TEXT)
    doc.add_paragraph()
    add_text(doc, "学位论文作者签名：__________________", size=FONT_XIAOSI, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(doc, "日期：2026年____月____日", size=FONT_XIAOSI, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_page_break(doc)
    add_heading_text(doc, "学位论文版权使用授权书", 1)
    add_text(doc, AUTHORIZATION_TEXT)
    doc.add_paragraph()
    add_text(doc, "本学位论文属于", size=FONT_XIAOSI, first_line_indent_cm=None)
    add_text(doc, "保  密□，在________年解密后适用本授权书。", size=FONT_XIAOSI, first_line_indent_cm=None)
    add_text(doc, "不保密■。", size=FONT_XIAOSI, first_line_indent_cm=None)
    doc.add_paragraph()
    add_text(doc, "学位论文作者签名：__________________    指导教师签名：__________________", size=FONT_XIAOSI, first_line_indent_cm=None)
    add_text(doc, "日期：2026年____月____日            日期：2026年____月____日", size=FONT_XIAOSI, first_line_indent_cm=None)


def add_abstracts(doc: Document) -> None:
    start_new_section(doc, None)
    add_text(doc, TITLE_CN, size=FONT_SAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="黑体")
    add_text(doc, "摘  要", size=FONT_SAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="黑体")
    for para in ABSTRACT_CN:
        add_text(doc, para)
    p = doc.add_paragraph()
    style_paragraph(p, first_line_indent_cm=None, space_before_pt=6)
    r1 = p.add_run("关键词：")
    set_run_font(r1, size=FONT_SI, bold=True, chinese="黑体")
    r2 = p.add_run(KEYWORDS_CN)
    set_run_font(r2, size=FONT_XIAOSI, chinese="宋体")

    start_new_section(doc, None)
    add_text(doc, TITLE_EN, size=FONT_SAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="Times New Roman", latin="Times New Roman")
    add_text(doc, "ABSTRACT", size=FONT_SAN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="Times New Roman", latin="Times New Roman")
    for para in ABSTRACT_EN:
        add_text(doc, para, chinese="Times New Roman", latin="Times New Roman")
    p = doc.add_paragraph()
    style_paragraph(p, first_line_indent_cm=None, space_before_pt=6)
    r1 = p.add_run("KEY WORDS: ")
    set_run_font(r1, size=FONT_SI, bold=True, chinese="Times New Roman", latin="Times New Roman")
    r2 = p.add_run(KEYWORDS_EN)
    set_run_font(r2, size=FONT_XIAOSI, chinese="Times New Roman", latin="Times New Roman")


def add_toc(doc: Document, tex: str) -> None:
    start_new_section(doc, None)
    add_heading_text(doc, "目  录", 1)
    outline = build_outline(tex)
    for level, text, bookmark_name in outline:
        p = doc.add_paragraph()
        style_paragraph(p, first_line_indent_cm=None)
        p.paragraph_format.left_indent = {1: Cm(0), 2: Cm(0.74), 3: Cm(1.48)}[level]
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(text)
        set_run_font(r, size=FONT_XIAOSI)
        p.add_run("\t")
        add_pageref_field(p, bookmark_name)


def parse_main_content(doc: Document, tex: str) -> None:
    body = extract_block(tex, r"\section{绪论}", r"\begin{thebibliography}{99}")
    body = r"\section{绪论}" + body
    lines = body.splitlines()
    i = 0
    sec = sub = subsub = 0
    fig_no = 0
    table_no = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith(r"\clearpage"):
            add_page_break(doc)
            i += 1
            continue

        if line.startswith(r"\section{"):
            sec += 1
            sub = 0
            subsub = 0
            fig_no = 0
            table_no = 0
            text = clean_inline(re.search(r"\\section\{(.*)\}", line).group(1))
            text = f"{sec} {text}"
            add_heading_text(doc, text, 1, f"toc_{sec}")
            i += 1
            continue

        if line.startswith(r"\subsection{"):
            sub += 1
            subsub = 0
            text = clean_inline(re.search(r"\\subsection\{(.*)\}", line).group(1))
            text = f"{sec}.{sub} {text}"
            add_heading_text(doc, text, 2, f"toc_{sec}_{sub}")
            i += 1
            continue

        if line.startswith(r"\subsubsection{"):
            subsub += 1
            text = clean_inline(re.search(r"\\subsubsection\{(.*)\}", line).group(1))
            text = f"{sec}.{sub}.{subsub} {text}"
            add_heading_text(doc, text, 3, f"toc_{sec}_{sub}_{subsub}")
            i += 1
            continue

        if line.startswith(r"\begin{enumerate}"):
            i += 1
            n = 1
            while i < len(lines) and not lines[i].strip().startswith(r"\end{enumerate}"):
                item_line = lines[i].strip()
                if item_line.startswith(r"\item"):
                    item_text = clean_inline(item_line[5:].strip())
                    add_text(doc, f"{n}. {item_text}", size=FONT_XIAOSI, first_line_indent_cm=None)
                    n += 1
                i += 1
            i += 1
            continue

        if line.startswith(r"\screenshotfig{") or line.startswith("screenshotfig{"):
            match = re.match(r"\\?screenshotfig\{([^{}]+)\}\{([^{}]+)\}\{([^{}]+)\}", line)
            image_name = clean_inline(match.group(1)) if match else ""
            caption = clean_inline(match.group(3)) if match else "截图占位"
            fig_no += 1
            add_image_or_placeholder(doc, image_name, f"图{sec}-{fig_no} {caption}")
            i += 1
            continue

        if line.startswith(r"\begin{figure}"):
            block_lines = [line]
            i += 1
            while i < len(lines):
                block_lines.append(lines[i])
                if lines[i].strip().startswith(r"\end{figure}"):
                    break
                i += 1
            block = "\n".join(block_lines)
            cap_match = re.search(r"\\caption\{([^{}]+)\}", block, re.S)
            caption = clean_inline(cap_match.group(1)) if cap_match else "图片占位"
            fig_no += 1
            add_placeholder(doc, caption)
            add_text(doc, f"图{sec}-{fig_no} {caption}", size=FONT_WU, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None, chinese="楷体")
            i += 1
            continue

        if line.startswith(r"\begin{table}") or line.startswith(r"\begin{longtable}"):
            block_lines = [line]
            end_tag = r"\end{table}" if line.startswith(r"\begin{table}") else r"\end{longtable}"
            i += 1
            while i < len(lines):
                block_lines.append(lines[i])
                if lines[i].strip().startswith(end_tag):
                    break
                i += 1
            caption, rows = parse_table_block("\n".join(block_lines))
            table_no += 1
            add_docx_table(doc, f"表{sec}-{table_no} {caption}", rows)
            i += 1
            continue

        if line.startswith("%") or line.startswith(r"\pagestyle") or line.startswith(r"\fancy") or line.startswith(r"\cfoot"):
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("\\") or nxt.startswith("%"):
                break
            para_lines.append(nxt)
            i += 1
        paragraph = clean_inline(" ".join(para_lines))
        if paragraph:
            add_text(doc, paragraph)


def add_references(doc: Document, tex: str) -> None:
    start_new_section(doc, "参考文献")
    add_heading_text(doc, "参考文献", 1)
    block = extract_block(tex, r"\begin{thebibliography}{99}", r"\end{thebibliography}")
    entries = re.split(r"\\bibitem\{[^{}]+\}", block)
    entries = [clean_inline(e) for e in entries if clean_inline(e)]
    for idx, entry in enumerate(entries, start=1):
        p = doc.add_paragraph()
        style_paragraph(p, first_line_indent_cm=None)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        run = p.add_run(f"[{idx}] {entry}")
        set_run_font(run, size=FONT_XIAOWU, chinese="宋体")


def add_acknowledgement(doc: Document) -> None:
    start_new_section(doc, None)
    add_heading_text(doc, "致谢", 1)
    paras = [
        "在本次毕业设计与论文撰写过程中，我得到了许多老师、同学和家人的帮助与支持。在此谨向所有关心、指导和帮助过我的人表示诚挚的感谢。",
        "首先，衷心感谢指导教师在选题确定、系统设计、论文撰写和细节修改过程中给予的耐心指导。老师严谨的治学态度和认真负责的工作作风，使我在项目开发与论文写作过程中受益匪浅。",
        "其次，感谢学院提供良好的学习与实验环境，感谢同学们在技术交流、系统测试和思路讨论方面给予的帮助。正是在不断交流与反复修改中，本课题才得以逐步完善。",
        "最后，感谢家人在学习期间给予的理解、鼓励与支持，使我能够顺利完成本次毕业设计。由于时间和个人能力所限，论文中仍难免存在不足，恳请各位老师批评指正。"
    ]
    for para in paras:
        add_text(doc, para)


def build_docx() -> Path:
    source_tex = next((p for p in SOURCE_CANDIDATES if p.exists()), None)
    if source_tex is None:
        raise FileNotFoundError("未找到可用的论文 tex 源文件。")
    tex = source_tex.read_text(encoding="utf-8")
    doc = Document()
    enable_update_fields_on_open(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    add_cover(doc)
    add_originality_and_authorization(doc)
    add_abstracts(doc)
    add_toc(doc, tex)
    start_new_section(doc, TITLE_CN)
    parse_main_content(doc, tex)
    add_references(doc, tex)
    add_acknowledgement(doc)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    if len(sys.argv) > 1:
        OUTPUT_DOCX = Path(sys.argv[1])
    output = build_docx()
    print(output)
